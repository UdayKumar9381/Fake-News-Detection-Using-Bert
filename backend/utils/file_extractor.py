"""
File Text Extraction Utilities
Extract text from PDF, DOCX, and Images
"""

import io
from typing import Optional
from PIL import Image, ImageOps, ImageFilter
import PyPDF2
import docx
import pytesseract


class FileExtractor:
    """Extract text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_bytes: PDF file as bytes
            
        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except Exception as e:
                    print(f"Warning: Could not extract page {page_num}: {e}")
                    continue
            
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"Error extracting PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_bytes: DOCX file as bytes
            
        Returns:
            Extracted text
        """
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_image(file_bytes: bytes, language: str = 'eng') -> str:
        """
        Extract text from image using OCR (Tesseract)
        
        Args:
            file_bytes: Image file as bytes
            language: OCR language (default: 'eng' for English)
            
        Returns:
            Extracted text
        """
        try:
            image = Image.open(io.BytesIO(file_bytes))

            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Basic preprocessing to improve OCR accuracy
            try:
                width, height = image.size
                if width < 1000:
                    new_width = 1600
                    new_height = int((new_width / width) * height)
                    image = image.resize((new_width, new_height), Image.LANCZOS)

                gray = ImageOps.grayscale(image)
                gray = ImageOps.autocontrast(gray)
                gray = gray.filter(ImageFilter.SHARPEN)
            except Exception as e:
                print(f"Warning: Image preprocessing failed: {e}")
                gray = image

            # Use Tesseract with a reasonable page segmentation mode
            try:
                config_str = '--psm 6'
                text = pytesseract.image_to_string(gray, lang=language, config=config_str)
                return text.strip()
            except Exception as e:
                print(f"Error: Tesseract OCR failed: {e}")
                # Check if tesseract is installed
                try:
                    import subprocess
                    subprocess.run(['tesseract', '--version'], capture_output=True)
                except FileNotFoundError:
                    return "[Error: Tesseract OCR engine not found on the system. Please install it to use image-to-text features.]"
                
                return f"[Error: Could not extract text from image: {str(e)}]"

        except Exception as e:
            print(f"Error opening image: {e}")
            raise ValueError(f"Error processing image file: {str(e)}")
    
    @staticmethod
    def get_file_info(file_bytes: bytes, filename: str) -> dict:
        """
        Get information about uploaded file
        
        Args:
            file_bytes: File as bytes
            filename: Original filename
            
        Returns:
            Dictionary with file information
        """
        file_size = len(file_bytes)
        file_extension = filename.split('.')[-1].lower() if '.' in filename else ''
        
        return {
            "filename": filename,
            "size_bytes": file_size,
            "size_mb": round(file_size / (1024 * 1024), 2),
            "extension": file_extension,
            "type": FileExtractor._get_file_type(file_extension)
        }
    
    @staticmethod
    def _get_file_type(extension: str) -> str:
        """Get file type from extension"""
        type_mapping = {
            'pdf': 'PDF Document',
            'docx': 'Word Document',
            'doc': 'Word Document',
            'jpg': 'Image',
            'jpeg': 'Image',
            'png': 'Image',
            'gif': 'Image'
        }
        return type_mapping.get(extension, 'Unknown')
    
    @staticmethod
    def is_supported_file(filename: str) -> bool:
        """Check if file type is supported"""
        supported_extensions = ['pdf', 'docx', 'doc', 'jpg', 'jpeg', 'png']
        extension = filename.split('.')[-1].lower() if '.' in filename else ''
        return extension in supported_extensions


def extract_text_from_file(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """
    Main function to extract text from any supported file type
    
    Args:
        file_bytes: File as bytes
        filename: Original filename
        
    Returns:
        Tuple of (extracted_text, file_info)
    """
    extractor = FileExtractor()
    
    # Check if file is supported
    if not extractor.is_supported_file(filename):
        raise ValueError(f"Unsupported file type: {filename}")
    
    # Get file info
    file_info = extractor.get_file_info(file_bytes, filename)
    
    # Extract text based on file type
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        text = extractor.extract_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        text = extractor.extract_from_docx(file_bytes)
    elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.gif')):
        text = extractor.extract_from_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
    
    return text, file_info


if __name__ == "__main__":
    # Test the file extractor
    print("File Extractor Module - Ready for use")
    print("Supported formats: PDF, DOCX, JPG, PNG")
    
    # Example usage:
    # with open('sample.pdf', 'rb') as f:
    #     file_bytes = f.read()
    #     text, info = extract_text_from_file(file_bytes, 'sample.pdf')
    #     print(f"Extracted {len(text)} characters")
    #     print(f"File info: {info}")